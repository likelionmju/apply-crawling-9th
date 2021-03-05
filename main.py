from selenium.webdriver import Chrome, ChromeOptions
from pathos.multiprocessing import ProcessPool
from pathlib import Path
from bs4 import BeautifulSoup
from requests import get
from zipfile import ZipFile
from time import time as current_time
from shutil import copyfileobj
from docx import Document


def is_sns(link) -> bool:
    for lk in sns_list:
        if lk in link:
            return True
    return False


def is_img(img: Path) -> bool:
    return img.suffix in img_extensions


def is_doc(doc: Path) -> bool:
    return doc.suffix in doc_extensions


def is_archive(archive: Path) -> bool:
    return archive.suffix in archive_extensions


def unzip(target: Path, to) -> None:
    with ZipFile(target) as zip_file:
        info = zip_file.infolist()
        for file in info:
            t = Path(file.filename)
            if t.is_dir():
                continue
            if not (is_doc(t) or is_img(t)):
                # if file is img or document continue
                continue
            if file.flag_bits != 2048:
                # if not utf-8
                file.filename = file.filename.encode("cp437").decode("cp949")
            zip_file.extract(file, to)


def reformat_file(file: Path) -> None:
    if is_img(file):
        file.rename(f"{file.parent}/시간표{file.suffix}")
    elif is_doc(file):
        file.rename(f"{file.parent}/포트폴리오{file.suffix}")


def download_file_by_url(u, save_path: Path) -> None:
    # source: https://stackoverflow.com/a/9419208
    r = get(u, stream=True)
    with open(save_path, 'wb') as fd:
        r.raw.decode_content = True
        copyfileobj(r.raw, fd)


def download_applicant_file(applicant: dict):
    if applicant["name"] in exclude_applicants:
        return
    with Path(f"./지원자 서류/{applicant['major']} {applicant['entrance_year'][2:]} {applicant['name']}") as path:
        if not path.exists():
            path.mkdir()
        if applicant["file"] == "X":
            return
        target_file = Path(f"{path}/{applicant['file'].split('/')[-1]}")
        download_file_by_url(applicant["file"], target_file)
        if is_archive(target_file):
            archive_dir = Path(f"{path}/시간표 및 포트폴리오")
            unzip(target_file, archive_dir)
            target_file.unlink()
        else:
            reformat_file(target_file)


def export_docx(applicant: dict):
    docx = Document()
    if applicant["name"] in exclude_applicants:
        return
    for key in applicant_ko_keys.keys():
        if "q" in key:
            docx.add_paragraph("")
            docx.add_paragraph(applicant_ko_keys[key]).bold = True
            docx.add_paragraph(applicant[key])
        else:
            docx.add_paragraph(f"{applicant_ko_keys[key]}: {applicant[key]}")
    docx.save(f"./지원자 서류/{applicant['major']} {applicant['entrance_year'][2:]} {applicant['name']}/지원서.docx")


class LikelionApplyCrawler:
    domain = "https://apply.likelion.org"
    apply_url = f"{domain}/apply"
    univ_url = f"{apply_url}/univ/"
    applicant_url = f"{apply_url}/applicant"

    __id_path = "id_username"
    __password_path = "id_password"
    __login_path = "//button[@type='submit']"
    __applicant_info_container_path = "#likelion_num"
    __answered_applicant_count_path = f"{__applicant_info_container_path} > div:nth-child(2) > p:nth-child(2)"
    __applicant_answer_container_path = ".answer_view > .applicant_detail_page"
    __applicants_path = f"{__applicant_info_container_path} > div.applicant_page > a"

    __html_parser = "html.parser"

    def __init__(self, admin_id: str, admin_pass: str, headless: bool = True) -> None:
        self.__admin_id = admin_id
        self.__admin_pass = admin_pass
        self.headless = headless
        self.univ_url += self.__admin_id.split('@')[0]
        self.applicant_pks = None
        print("crawling start...")

    def __enter__(self):
        self.cookies = {ck["name"]: ck["value"] for ck in self.login()}
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        print(f"async multi processing time: {current_time() - start_time}s")

    def login(self) -> dict:
        driver_options = ChromeOptions()
        driver_options.headless = self.headless
        with Chrome(executable_path="./chromedriver", options=driver_options) as driver:
            driver.get(self.apply_url)
            driver.find_element_by_id(self.__id_path).send_keys(self.__admin_id)
            driver.find_element_by_id(self.__password_path).send_keys(self.__admin_pass)
            driver.find_element_by_xpath(self.__login_path).submit()
            print(f"finish login time: {current_time() - start_time}s")
            return driver.get_cookies()

    def request_univ_page_source(self) -> str:
        return get(self.univ_url, cookies=self.cookies).text

    #
    # def get_answered_applicant_count(self, univ_page_source: str) -> int:
    #     univ_page = bs.BeautifulSoup(univ_page_source, features=self.__html_parser)
    #     info_text = univ_page.select_one(self.__answered_applicant_count_path).string
    #     find = self.__applicant_count_regex.search(info_text).group()
    #     return int(find)

    def extract_all_applicant_pks(self, univ_page_source: str) -> None:
        univ_page = BeautifulSoup(univ_page_source, features=self.__html_parser)
        self.applicant_pks = [applicant.get("href").split("/")[-1]
                              for applicant in univ_page.select(self.__applicants_path)]

    def request_applicant_source(self, applicant_pk: str) -> str:
        return get(f"{self.applicant_url}/{applicant_pk}", cookies=self.cookies).text

    def parse_applicant_page(self, page) -> dict:
        try:
            soup = BeautifulSoup(page, features="html.parser")
            applicant_info_container = soup.select_one(self.__applicant_info_container_path)
            applicant_answer_container = soup.select_one(self.__applicant_answer_container_path)

            applicant_name = applicant_info_container.find("h3").string
            if applicant_name in exclude_applicants:
                return {"name": applicant_name}
            user_info_list = applicant_info_container.select("div.row")
            user_answer_list = applicant_answer_container.select("div.m_mt")
            additional = [user_info.contents[1].get("href") for user_info in user_info_list[2:]]
            git = sns = applicant_file = None

            for item in additional:
                if item is None:
                    continue
                if "git" in item:
                    git = item
                elif is_sns(item):
                    sns = item
                elif "cdn" in item:
                    applicant_file = item

            applicant = {
                "name": applicant_name,
                "entrance_year": user_info_list[0].contents[1].text,
                "major": user_info_list[0].contents[-2].text,
                "phone_num": user_info_list[1].contents[1].text,
                "email": user_info_list[1].contents[-2].text,
                "git": git if git is not None else "X",
                "sns": sns if sns is not None else "X",
                "file": applicant_file if applicant_file is not None else "X",
                "q1": user_answer_list[0].contents[1].text,
                "q2": user_answer_list[1].contents[1].text,
                "q3": user_answer_list[2].contents[1].text,
                "q4": user_answer_list[3].contents[1].text,
                "q5": user_answer_list[4].contents[1].text,
            }
            phone_num = applicant["phone_num"]
            if len(phone_num) == 11:
                formatting_phone_num = [phone_num[:3], phone_num[3:7], phone_num[7:]]
                applicant["phone_num"] = "-".join(formatting_phone_num)
            return applicant
        except AttributeError:
            print("error")

    # def export_csv(self) -> None:
    #     with open("지원자목록.csv", "w", newline="", encoding="utf-8") as file:
    #         keys_to_header = {
    #             "name": "이름",
    #             "entrance_year": "입학 년도",
    #             "major": "전공",
    #             "phone_num": "전화번호",
    #             "email": "이메일",
    #             "git": "github",
    #             "sns": "SNS",
    #             "file": "file",
    #             "q1": "지원 동기",
    #             "q2": "만들고 싶은 서비스",
    #             "q3": "가장 기억에 남는 활동과 느낀 점",
    #             "q4": "기억에 남는 프로그래밍 경험과 느낀 점 또는 배우고 싶은 것",
    #             "q5": "1학기 시간표와 opt(포트폴리오)"
    #         }
    #         writer = csv.DictWriter(file, fieldnames=keys_to_header.keys())
    #         writer.writerow(keys_to_header)
    #         for applicant in self.applicants.values():
    #             writer.writerow(applicant)


if __name__ == "__main__":
    import sys
    sys.setrecursionlimit(3000)
    required_dir = Path("./지원자 서류")
    if not required_dir.exists():
        required_dir.mkdir()
    sns_list = ("facebook", "instagram", "twitter")
    img_extensions = (".png", ".jpg", ".jpeg", ".PNG", ".JPG", ".JPEG")
    doc_extensions = (".pdf", ".docx", ".hwp")
    archive_extensions = (".zip", ".tar.gz", ".rar", ".7z")
    a_id = input("관리자 ID: ")
    a_pass = input("관리자 PW: ")
    exclude_applicants = input("제외할 사람: ").split() or ["테스트", "한준혁", "김예빈", "박성제"]
    applicant_ko_keys = {
        "name": "이름",
        "entrance_year": "입학 년도",
        "major": "전공",
        "phone_num": "전화번호",
        "email": "이메일",
        "git": "GitHub",
        "sns": "SNS",
        "q1": "지원 동기",
        "q2": "만들고 싶은 서비스",
        "q3": "참여했던 팀 활동 중 가장 기억에 남는 활동과 느낀 점을 작성해주세요.",
        "q4": "기억에 남는 프로그래밍 경험과 느낀 점을 작성해주세요. 만약 없다면, 어떤 것을 배우고 싶은지 작성해주세요.",
        "q5": "첨부파일에 1학기 시간표를 캡쳐해서 제출해주세요. 만약 포트폴리오를 제출하고 싶으시다면 함께 제출해주세요. (압축파일로 제출해주시면 됩니다.)",
    }

    start_time = current_time()
    with LikelionApplyCrawler(admin_id=a_id, admin_pass=a_pass) as c:
        source = c.request_univ_page_source()
        c.extract_all_applicant_pks(source)
        print(f"finish extract pks time: {current_time() - start_time}s")
        with ProcessPool() as main_pool:
            applicant_sources = main_pool.amap(c.request_applicant_source, c.applicant_pks).get()
            print(f"finish request sources time: {current_time() - start_time}s")
            applicants = main_pool.amap(c.parse_applicant_page, applicant_sources).get()
            print(f"total applicants count: {len(applicants)}")
            main_pool.map(download_applicant_file, applicants)
            print(f"finish download time: {current_time() - start_time}s")
            main_pool.map(export_docx, applicants)
